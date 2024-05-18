from langchain.document_loaders import PyPDFLoader
from utils import *
from prompts import *
import os
from langchain_google_genai import ChatGoogleGenerativeAI
from dotenv import load_dotenv
import streamlit as st
from docx import Document
import io
import base64
import firebase_admin
from firebase_admin import auth, exceptions, credentials, initialize_app
from google.oauth2 import id_token
from google.auth.transport import requests
from firebase_admin import credentials

# Define the service account key as a dictionary
service_account_key = {
  "type": "service_account",
  "project_id": "web-development-631c2",
  "private_key_id": "458945d066d573f9becd93bf0fb57d070c1f4ae7",
  "private_key": "-----BEGIN PRIVATE KEY-----\nMIIEvwIBADANBgkqhkiG9w0BAQEFAASCBKkwggSlAgEAAoIBAQCQ9UHe+U1viRUy\nerdm2NXYEeN3xlG4cphXoxT3NZ8SIKzx1pEcERbPIDRviP3M9Kux6gQQPKdnB8k2\nIrJ1iOjFoUspNIvZb6Wy39t1eSUwdTDC4XkqFU4i8pPIMY3U0xFbIVZ+tooNx3HD\nLKnuTCr6PAA9mzxBnn9DEZOlko1Mg+G100znCgK53ofCx6B3C6T/7o2b3eul2Hq3\nJcEZO7zXkpp7Mc96fDFldjqypEkQkk/Hrc34Y3cohpBVlbtZuRIoTZqxYa6glnzR\ntmc0LwryKQJOebCGJLnXIXcuoOcs/kkNmxBkD5KKbgIIOmwD491ESdkwJmoB5Ld0\nusPBLPWDAgMBAAECggEAJqoMSlWT1nqpeT/rOFcA3vOSiSmQNZL3WHTKzE90oOiv\nSIgzR1m1jESqR8+x5p+iRj/FmrEa86vLexGRKTF96OznN9Pos3QIrefwyOopxb2U\nZn6ui0+IemDakyzIiGpN+FExmXXF5/v5N5ZrswCE/B7R75y+2z/T8Qc4qxKmdrBV\nwbnAozP7FyCfeBogz0FEUD5CjTXMstkn3/hP1dvCUtPg9Bw9pO6iOA/XYnYDof0g\nE2FPUvAEckgxNMuu9QnOQuCn9Bz76fvE9TtZP/FIbCYUWmNgoCj1cxH7oSrtXhDr\n6kYEzAaMqij+hGsLlKa/X9p5+R72cNXf9ZvWCi5LcQKBgQDJA19QGad4HFc06EqP\nQtNKaP5yG65xD2Sm8GG9Si7E1yjW7o67lOCTgUe52LpxIErtb9TDFRqS417eMUUS\nM98UMnARBGbndSORFi3vTBFOhU/jw99uNZCkokw2+guhFOpGsyTtzsQoWx26DZpw\n05/Im6tykoUUJS1OeupPAtRqswKBgQC4nG8xfOjWj+m+Q1vy9siggvByf8cP56N2\nvFBWdklectwohqqYGeLVKIVqCZHEjHhbEIdk74XaTMaUFVmApsgC08BWnqLGxlhg\nuSYXc4F4PzURNSzxKrhdK8LxF2yee78+rh0yFYrJqptKKZF19+o4ImpIP4XpdNIr\nVfLCbzbx8QKBgQCfYFcf8VdA+meCgXMsZ3WTl++2lZAYAHZzjkBQjxLJ6zZsQ7VT\nsVTmXyh5iAc7TrlZzMNK8d6h/DjpvvLe3dXP87W7KFxWxB+xF37QLuXB9h1Dn7f1\nb+8BOdVn+ZVi4GxoJf3vt3L79FJ7g/Gh54jWy+aB7v4o3NkxUm+Yucul1QKBgQCj\ng+v5KWoSoeWnYIJQ8cVCZrgSspreu41N3f5KlvrlbVcTjjQMBOGla/6/rTovnTv9\nbOZ/wKgOc3JwN3jPOuptlqEGF/yVk5k0EspqeyRyHYojc7Ya5nvmCpzgbp8GwRTj\nr9SsR+hmN3bUeX0dwDj0/4aj95/k6FQ/jiE3lspUAQKBgQDEGe7LeUWBHdSiyDhR\nIbA8/XYSRpq2z+Ha38Cgwgq3K660cO9s8EPpgCSWI6J8Hc8lAHCnR57qNKOj6qbm\nJo6Vy0W9MK7MH5AewqZHXVCrobMsMGLkId9Nzjbt3Dpdfk3icrlsOCy3ukLtJ2Ul\npvN5EMluI64KUjgX9H5pSk2lTw==\n-----END PRIVATE KEY-----\n",
  "client_email": "firebase-adminsdk-m16gy@web-development-631c2.iam.gserviceaccount.com",
  "client_id": "107463406523932663440",
  "auth_uri": "https://accounts.google.com/o/oauth2/auth",
  "token_uri": "https://oauth2.googleapis.com/token",
  "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
  "client_x509_cert_url": "https://www.googleapis.com/robot/v1/metadata/x509/firebase-adminsdk-m16gy%40web-development-631c2.iam.gserviceaccount.com",
  "universe_domain": "googleapis.com"
}

# Initialize the app with a service account, granting admin privileges
cred = credentials.Certificate(service_account_key)
try:
    firebase_admin.get_app()
except ValueError as e:
    initialize_app(cred)

st.set_page_config(layout="wide", page_title="MCQ Generation from Documents", page_icon='deep-learning.png')

temperature = 0.3
pages = []
numPairs = 2

load_dotenv()

google_api_key = os.environ.get("AIzaSyB3ftUCJdTJYCc6pXJm3lxGLdv0hYZE3_U")

# Create ChatGoogleGenerativeAI object with the API key
model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.3, google_api_key="AIzaSyB3ftUCJdTJYCc6pXJm3lxGLdv0hYZE3_U")


# Function to extract text from a Word document
def extract_text_from_word(uploaded_file):
    text = ""
    docx_file = Document(uploaded_file)
    for paragraph in docx_file.paragraphs:
        text += paragraph.text + "\n"
    return text

# Function to generate MCQ question-answer pairs
def McqQAPairs(context, numPairs, model):
    mcq_string = ""
    with st.spinner('Generating MCQ Question Answer Pairs...'):
        response = getMcqQAPairs(context, numPairs, model)

    if response is None:
        st.error('Failed to generate MCQ Question Answer Pairs.')
        return None

    for i, qaPair in enumerate(response):
        with st.chat_message("user"):
            st.write(f"{i+1}) {qaPair['question']}")
            mcq_string += f"{str(i+1)}){qaPair['question']}\n"
            for j, option in enumerate(qaPair['options']):
                st.write(f"{string.ascii_uppercase[j]}) {option}")
                mcq_string += f"{string.ascii_uppercase[j]}) {option}\n"
            st.write(f"Answer: {qaPair['options'][qaPair['correct_option_index']]}")
            mcq_string += f"Answer: {qaPair['options'][qaPair['correct_option_index']]}\n\n"

    return mcq_string

# Function to render the MCQ Generator page
def mcq_generator():
    if not st.session_state.get("logged_in"):
        st.error("You need to login first.")
        return 
    st.title('MCQ Generation From Documents')

    mcq_string = None

    with st.container(border=True):
        col1, col2 = st.columns([2, 1])
        with col1:
            st.write("Please Upload Your File")
            uploaded_file = st.file_uploader("Choose a file", type=['.pdf', '.docx'], accept_multiple_files=False)
            if uploaded_file is not None:
                if uploaded_file.type == "application/pdf":
                    with open("temp.pdf", "wb") as f:
                        f.write(uploaded_file.getbuffer())
                    file_path = "temp.pdf"
                    pdf_loader = PyPDFLoader(file_path)
                    pages = pdf_loader.load_and_split()
                elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    text = extract_text_from_word(uploaded_file)
                    pages = [text]

        with col2:
            numPairs = st.number_input('Number of QA Pairs', min_value=1, max_value=20, step=2, value=2)
            optionCategory = ("PDF", "Word")
            document_type = st.selectbox("Choose Document Type:", optionCategory)

    if st.button("Generate MCQs"):  
        if len(pages) and uploaded_file is not None:
            mcq_string = McqQAPairs(pages, numPairs, model)  
            if mcq_string:  
                # Store the generated MCQ string in the Streamlit session state
                st.session_state.mcq_string = mcq_string
            else:
                st.error('Failed to generate MCQs.')
        else:
            st.error('Required Fields are Missing!', icon="🚨")

# Function to render the Download page
def download_page():
    if not st.session_state.get("logged_in"):
        st.error("You need to login first.")
        return
    st.title("Download Generated MCQs")
    
    st.write("Click the button below to download the generated MCQs.")
    
    mcq_string = st.session_state.get("mcq_string", "")
    if mcq_string:
        document = Document()
        document.add_paragraph(mcq_string)
        
        document_stream = io.BytesIO()
        document.save(document_stream)
        document_stream.seek(0)
        b64 = base64.b64encode(document_stream.read()).decode()
        href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="generated_mcqs.docx">Download MCQs</a>'
        st.markdown(href, unsafe_allow_html=True)
    else:
        st.error("No MCQs found.")

# Function to render the Home page
def home_page():
    if not st.session_state.get("logged_in"):
        st.error("You need to login first.")
        return
    st.title("MCQ Generator App")

    st.write("## App Overview:")
    st.markdown("<p style='font-size:20px;'>The MCQ Generator App presents a convenient solution for automating the creation of multiple-choice questions (MCQs) from your text materials. Whether you're an educator developing quizzes, a content producer designing evaluations, or someone in need of prompt and effective MCQ generation, this application is your go-to option!</p>", unsafe_allow_html=True)
    
    st.write("## Main Features:")
    st.markdown("<p style='font-size:20px;'>- <strong>Document Upload:</strong> Easily upload PDF or Word files containing your text material.</p>", unsafe_allow_html=True)
    st.markdown("<p style='font-size:20px;'>- <strong>MCQ Generation:</strong> Instantly produce MCQs aligned with the content of your uploaded documents.</p>", unsafe_allow_html=True)
    st.markdown("<p style='font-size:20px;'>- <strong>Customization:</strong> Specify the number of MCQ pairs questions which you want to generate.</p>", unsafe_allow_html=True)
    st.markdown("<p style='font-size:20px;'>- <strong>Download:</strong> Effortlessly retrieve the generated MCQs in a Word document for your convenience.</p>", unsafe_allow_html=True)

     
    st.write("## Getting Started:")
    st.markdown("<p style='font-size:20px;'>1. Navigate to the <strong>MCQ Generator</strong> section through the sidebar menu.</p>", unsafe_allow_html=True)
    st.markdown("<p style='font-size:20px;'>2. Upload your document and indicate the desired number of MCQ pairs.</p>", unsafe_allow_html=True)
    st.markdown("<p style='font-size:20px;'>3. Click on <strong>Generate MCQs</strong> to create MCQs based on your document.</p>", unsafe_allow_html=True)
    st.markdown("<p style='font-size:20px;'>4. Review the generated MCQs and proceed to the <strong>Download</strong> page to grab them in a Word document.</p>", unsafe_allow_html=True)
    
    st.write("## Initiate MCQ Generation Now!")
    st.markdown("<p style='font-size:20px;'>Explore the sidebar menu and commence effortlessly generating MCQs!</p>", unsafe_allow_html=True)


def login_signup_page():
    st.title("Login/Sign-up Page")
    
    # Display login/sign-up form
    email = st.text_input("Email")
    password = st.text_input("Password", type="password")
    login_btn = st.button("Login")
    signup_btn = st.button("Sign Up")
    
    if signup_btn:
        # Perform sign-up process
        try:
            user = auth.create_user(email=email, password=password)
            st.success("Sign-up successful! You can now login.")
        except exceptions.FirebaseError as e:
            st.error("Sign-up failed. Please try again.")

    if login_btn:
        # Perform login authentication
        try:
            user = auth.get_user_by_email(email) # Corrected method name
            st.session_state.email = email
            st.session_state.password = password
            st.session_state.logged_in = True     
        except exceptions.FirebaseError as e:
            st.error("Login failed. Please check your credentials.")
    
    if st.session_state.get("logged_in"):
        # Change the color of the text and remove link behavior
        st.write(f'<span style="color: black; text-decoration: none;">Logged in as: {st.session_state.email}</span>', unsafe_allow_html=True)


# Create a dictionary to map page names to corresponding functions
pages = {
    "Acess Portal": login_signup_page,
    "Home": home_page,
    "MCQ Generator": mcq_generator,
    "Download": download_page
}

# Render the sidebar menu for page selection
selected_page = st.sidebar.selectbox("Select Page", list(pages.keys()))

# Call the selected page function to render its content
pages[selected_page]()

